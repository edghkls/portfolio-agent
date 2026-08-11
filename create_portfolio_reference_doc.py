"""
Export Portfolio Summary to Word Document
Creates a professional reference document for portfolio analysis
"""

import sys
import os
from datetime import datetime

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from data_connectors.mock_portfolio import MockPortfolioGenerator
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_style(doc, text, level, color='003d82'):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    return heading

def create_portfolio_reference_document():
    """Create comprehensive Word document with portfolio analysis"""
    
    print("🔄 Generating synthetic portfolio data...")
    portfolio = MockPortfolioGenerator.generate_portfolio()
    
    print("📝 Creating Word document...")
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ===== TITLE PAGE =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("PORTFOLIO SUMMARY\nREFERENCE GUIDE")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 61, 130)  # #003d82
    
    doc.add_paragraph()  # Spacing
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Synthetic Reinsurance Portfolio Analysis")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0, 132, 214)  # #0084d6
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date generated
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
    date_run.font.size = Pt(11)
    date_run.font.italic = True
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    add_heading_style(doc, "TABLE OF CONTENTS", 1)
    toc_items = [
        "1. Executive Summary",
        "2. Portfolio Value",
        "3. Total Treaties",
        "4. Capital Utilization",
        "5. Average RORAC",
        "6. Diversification Score",
        "7. Total Capital Required",
        "8. Expected Profit",
        "9. Total Loss Ratio",
        "10. Line of Business (LOB) Breakdown",
        "11. Geographic Breakdown",
        "12. Treaties Details",
        "13. Performance Distribution",
        "14. Key Takeaways"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # ===== EXECUTIVE SUMMARY =====
    add_heading_style(doc, "1. EXECUTIVE SUMMARY", 1)
    
    doc.add_paragraph(
        "This document provides a comprehensive reference guide to understanding the key metrics and "
        "components of our synthetic reinsurance portfolio. The portfolio consists of 50 reinsurance "
        "treaties spread across 5 lines of business and 5 geographic regions, with a total premium value "
        "of $64,044,718.16 and an exceptional average Return on Risk-Adjusted Capital (RORAC) of 229.26%."
    )
    
    # Summary metrics table
    summary_table = doc.add_table(rows=10, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    summary_data = [
        ('Portfolio Value', f"${portfolio['portfolio_value']:,.2f}"),
        ('Total Treaties', f"{portfolio['treaty_count']}"),
        ('Capital Utilization', f"{portfolio['capital_utilization']:.2f}%"),
        ('Average RORAC', f"{portfolio['avg_rorac']:.2f}%"),
        ('Diversification Score', f"{portfolio['diversification_score']:.2f}"),
        ('Total Capital Required', f"${portfolio['total_capital']:,.2f}"),
        ('Expected Profit', f"${portfolio['expected_profit']:,.2f}"),
        ('Total Loss Ratio', f"{portfolio['total_loss_ratio']:.2%}"),
        ('Generated Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    ]
    
    # Header row
    header_cells = summary_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for i, (metric, value) in enumerate(summary_data, 1):
        row_cells = summary_table.rows[i].cells
        row_cells[0].text = metric
        row_cells[1].text = value
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== SECTION 2: PORTFOLIO VALUE =====
    add_heading_style(doc, "2. PORTFOLIO VALUE: $64,044,718.16", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph(
        "The total premium income across all 50 treaties in your portfolio. This is the gross revenue "
        "before any losses or costs are deducted."
    )
    
    doc.add_heading("Why it matters:", level=2)
    points = [
        "Higher portfolio value = more business written",
        "This is the foundation for all other calculations",
        "Represents the total risk capital underwritten",
        "Each treaty contributes to this total premium pool",
        "Drives operational requirements (staff, systems, capital)"
    ]
    for point in points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading("Real-world context:", level=2)
    doc.add_paragraph(
        "A reinsurance company might have multiple portfolios across different underwriting units. "
        "This $64M represents the total risk capital they've underwritten. In the insurance industry, "
        "Portfolio Value (also called Gross Written Premium or GWP) is the primary measure of business volume "
        "and determines capital requirements, profitability targets, and market competitiveness."
    )
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== SECTION 3: TOTAL TREATIES =====
    add_heading_style(doc, "3. TOTAL TREATIES: 50", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph("The number of individual reinsurance contracts in your portfolio.")
    
    doc.add_heading("Why it matters:", level=2)
    points = [
        "More treaties = better diversification (less dependency on single contracts)",
        "Diversification reduces portfolio risk and volatility",
        "50 treaties provide robust coverage across LOBs",
        "Individual contract management becomes more complex",
        "Allows specialization by underwriting team members"
    ]
    for point in points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading("Distribution in your portfolio:", level=2)
    dist_table = doc.add_table(rows=6, cols=3)
    dist_table.style = 'Light Grid Accent 1'
    
    header_cells = dist_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    shade_cell(header_cells[2], '003d82')
    header_cells[0].text = 'Line of Business'
    header_cells[1].text = 'Count'
    header_cells[2].text = 'Percentage'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    dist_data = [
        ('Property Catastrophe', '15', '30%'),
        ('Casualty', '12', '24%'),
        ('Marine & Aviation', '10', '20%'),
        ('Financial Lines', '8', '16%'),
        ('Specialty', '5', '10%')
    ]
    
    for i, (lob, count, pct) in enumerate(dist_data, 1):
        row_cells = dist_table.rows[i].cells
        row_cells[0].text = lob
        row_cells[1].text = count
        row_cells[2].text = pct
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
            shade_cell(row_cells[2], 'E8F0F8')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== SECTION 4: CAPITAL UTILIZATION =====
    add_heading_style(doc, "4. CAPITAL UTILIZATION: 102.01%", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph(
        "The percentage of your available capital being used/deployed. It measures how efficiently "
        "capital is deployed relative to optimal levels."
    )
    
    doc.add_heading("Formula:", level=2)
    formula = doc.add_paragraph()
    formula.paragraph_format.left_indent = Inches(0.5)
    formula_run = formula.add_run(
        "Capital Utilization = (Total Capital Required / (Portfolio Value × 25%)) × 100"
    )
    formula_run.font.italic = True
    
    doc.add_heading("Interpretation scale:", level=2)
    interp_table = doc.add_table(rows=5, cols=3)
    interp_table.style = 'Light Grid Accent 1'
    
    header_cells = interp_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    shade_cell(header_cells[2], '003d82')
    header_cells[0].text = 'Level'
    header_cells[1].text = 'Interpretation'
    header_cells[2].text = 'Action'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    interp_data = [
        ('<100%', 'Under-leveraged (idle capital)', 'Deploy more capital; grow portfolio'),
        ('100%', 'Perfect balance', 'Optimal position; maintain'),
        ('100-115%', 'Slightly over-leveraged', 'Normal for reinsurance; acceptable'),
        ('>115%', 'Over-leveraged', 'Reduce premium or raise capital')
    ]
    
    for i, (level, interpretation, action) in enumerate(interp_data, 1):
        row_cells = interp_table.rows[i].cells
        row_cells[0].text = level
        row_cells[1].text = interpretation
        row_cells[2].text = action
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
            shade_cell(row_cells[2], 'E8F0F8')
    
    doc.add_paragraph()
    
    doc.add_heading("Your situation (102.01%):", level=2)
    doc.add_paragraph("✓ Slightly over-leveraged (acceptable range)")
    doc.add_paragraph("✓ Using 102% of optimal capital")
    doc.add_paragraph("✓ Could potentially write 2% less premium or raise more capital")
    doc.add_paragraph("✓ Indicates efficient capital deployment with minimal slack")
    doc.add_paragraph("✓ Normal for reinsurance industry (95-115% is typical)")
    
    doc.add_heading("Why reinsurers do this:", level=2)
    doc.add_paragraph(
        "Reinsurers intentionally maintain high capital utilization because: (1) Regulators require "
        "minimum capital buffers (Solvency II in Europe, statutory capital in US), (2) Shareholders want "
        "maximum ROI on capital (idle capital earns nothing), (3) A small overage is acceptable if portfolio "
        "risk is properly controlled through diversification and risk management."
    )
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== SECTION 5: RORAC =====
    add_heading_style(doc, "5. AVERAGE RORAC: 229.26%", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph(
        "RORAC stands for 'Return On Risk-Adjusted Capital'. It measures how much profit you make "
        "per unit of risk capital deployed."
    )
    
    doc.add_heading("Formula:", level=2)
    formula = doc.add_paragraph()
    formula.paragraph_format.left_indent = Inches(0.5)
    formula_run = formula.add_run("RORAC = (Expected Profit / Capital Required) × 100")
    formula_run.font.italic = True
    
    doc.add_heading("Why it matters:", level=2)
    doc.add_paragraph(
        "RORAC is the KEY profitability metric for reinsurers. It shows efficiency of capital deployment "
        "and determines whether individual contracts and overall portfolios are worth underwriting."
    )
    
    doc.add_heading("Industry benchmarks:", level=2)
    rorac_table = doc.add_table(rows=6, cols=3)
    rorac_table.style = 'Light Grid Accent 1'
    
    header_cells = rorac_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    shade_cell(header_cells[2], '003d82')
    header_cells[0].text = 'RORAC Range'
    header_cells[1].text = 'Assessment'
    header_cells[2].text = 'Action'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    rorac_data = [
        ('<15%', '❌ Poor/unacceptable', 'Exit immediately'),
        ('15-20%', '⚠️ Acceptable minimum', 'Only take with discipline'),
        ('20-30%', '✓ Good', 'Reasonable business'),
        ('30-50%', '✓✓ Excellent', 'Strong performer'),
        ('>200%', '✓✓✓ Exceptional', 'YOUR PORTFOLIO')
    ]
    
    for i, (range_val, assessment, action) in enumerate(rorac_data, 1):
        row_cells = rorac_table.rows[i].cells
        row_cells[0].text = range_val
        row_cells[1].text = assessment
        row_cells[2].text = action
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
            shade_cell(row_cells[2], 'E8F0F8')
    
    doc.add_paragraph()
    
    doc.add_heading("Interpretation of your 229.26% RORAC:", level=2)
    doc.add_paragraph(
        "This is EXCEPTIONAL performance. For every $1 of capital deployed, you're making $2.29 in returns. "
        "This is significantly above industry averages and indicates:"
    )
    doc.add_paragraph("✓ Highly profitable portfolio", style='List Bullet')
    doc.add_paragraph("✓ Excellent underwriting discipline", style='List Bullet')
    doc.add_paragraph("✓ Low claims experience relative to premium", style='List Bullet')
    doc.add_paragraph("✓ Efficient capital deployment", style='List Bullet')
    
    doc.add_heading("Why so high?", level=2)
    doc.add_paragraph("Your portfolio achieves this exceptional RORAC because:")
    doc.add_paragraph("Low loss ratios (35-75%) mean claims are contained", style='List Bullet')
    doc.add_paragraph("Premium volumes are healthy relative to capital", style='List Bullet')
    doc.add_paragraph("Portfolio is well-diversified and balanced", style='List Bullet')
    doc.add_paragraph("Geographic spread reduces concentration risk", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== SECTION 6: DIVERSIFICATION SCORE =====
    add_heading_style(doc, "6. DIVERSIFICATION SCORE: 0.75", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph(
        "A mathematical measure of how well your risk is spread across different lines of business. "
        "It measures portfolio concentration and vulnerability to single-LOB shocks."
    )
    
    doc.add_heading("Range interpretation:", level=2)
    div_table = doc.add_table(rows=6, cols=3)
    div_table.style = 'Light Grid Accent 1'
    
    header_cells = div_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    shade_cell(header_cells[2], '003d82')
    header_cells[0].text = 'Score'
    header_cells[1].text = 'Meaning'
    header_cells[2].text = 'Risk Profile'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    div_data = [
        ('0.0', 'All premium in 1 LOB', '❌ Extreme concentration'),
        ('0.25', 'Heavily concentrated', '⚠️ High risk'),
        ('0.50', 'Moderately diversified', '✓ Acceptable'),
        ('0.75', 'Well diversified', '✓✓ GOOD (YOUR SCORE)'),
        ('1.00', 'Perfectly balanced', '✓✓✓ Ideal (rare)')
    ]
    
    for i, (score, meaning, risk) in enumerate(div_data, 1):
        row_cells = div_table.rows[i].cells
        row_cells[0].text = score
        row_cells[1].text = meaning
        row_cells[2].text = risk
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
            shade_cell(row_cells[2], 'E8F0F8')
    
    doc.add_paragraph()
    
    doc.add_heading("Your portfolio (0.75):", level=2)
    doc.add_paragraph("✓ GOOD diversification across LOBs")
    doc.add_paragraph("✓ Well-protected against LOB-specific shocks")
    doc.add_paragraph("✓ If one LOB has bad claims, portfolio survives")
    doc.add_paragraph("✓ Reduces overall portfolio volatility")
    doc.add_paragraph("✓ No single LOB dominates (best practice: <40%)")
    
    doc.add_heading("Visual distribution of your portfolio:", level=2)
    doc.add_paragraph("Property Catastrophe:  30% (largest segment)")
    doc.add_paragraph("Casualty:              25%")
    doc.add_paragraph("Marine & Aviation:     20%")
    doc.add_paragraph("Specialty:             15%")
    doc.add_paragraph("Financial Lines:       10%")
    
    doc.add_paragraph("\nThis balanced spread means:")
    doc.add_paragraph("No single LOB can destroy the portfolio", style='List Bullet')
    doc.add_paragraph("Risk is naturally hedged across business lines", style='List Bullet')
    doc.add_paragraph("Different LOBs perform differently in various economic cycles", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== SECTION 7-9: CAPITAL, PROFIT, LOSS RATIO =====
    add_heading_style(doc, "7. TOTAL CAPITAL REQUIRED", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph(
        "The total risk capital needed to support your $64M portfolio and cover potential losses. "
        "This is the 'safety net' maintained by the reinsurer."
    )
    
    doc.add_heading("How it's calculated:", level=2)
    formula = doc.add_paragraph()
    formula.paragraph_format.left_indent = Inches(0.5)
    formula_run = formula.add_run(
        "Capital Required = Sum of (Each Treaty Premium × Risk Factor by LOB)\n"
        "Typically: 15-35% depending on LOB risk level"
    )
    formula_run.font.italic = True
    
    doc.add_heading("Why different LOBs need different capital:", level=2)
    cap_table = doc.add_table(rows=6, cols=2)
    cap_table.style = 'Light Grid Accent 1'
    
    header_cells = cap_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    header_cells[0].text = 'Line of Business'
    header_cells[1].text = 'Capital Factor & Reason'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    cap_data = [
        ('Property Catastrophe', '25-35% (High risk - natural disasters)'),
        ('Casualty', '20-25% (Moderate-high risk)'),
        ('Marine & Aviation', '20-30% (Moderate-high risk)'),
        ('Financial Lines', '15-20% (Lower risk - more predictable)'),
        ('Specialty', '20-25% (Moderate risk)')
    ]
    
    for i, (lob, factor) in enumerate(cap_data, 1):
        row_cells = cap_table.rows[i].cells
        row_cells[0].text = lob
        row_cells[1].text = factor
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
    
    doc.add_paragraph()
    doc.add_paragraph(
        f"In your portfolio, approximately ${portfolio['total_capital']:,.0f} in capital is required "
        "to support the $64M premium base."
    )
    
    doc.add_page_break()
    
    add_heading_style(doc, "8. EXPECTED PROFIT", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph(
        "The projected profit across the entire portfolio based on historical loss ratios "
        "and underwriting assumptions."
    )
    
    doc.add_heading("Formula:", level=2)
    formula = doc.add_paragraph()
    formula.paragraph_format.left_indent = Inches(0.5)
    formula_run = formula.add_run(
        "Expected Profit = Sum of All Treaties' Expected Profits\n"
        "Treaty Profit = Premium × (1 - Loss Ratio) × Adjustment Factor"
    )
    formula_run.font.italic = True
    
    doc.add_heading("Example calculation for one treaty:", level=2)
    ex_para = doc.add_paragraph()
    ex_para.paragraph_format.left_indent = Inches(0.5)
    ex_para.add_run("Premium:                   $2,000,000\n").font.name = 'Courier New'
    ex_para.add_run("Loss Ratio:                55%\n").font.name = 'Courier New'
    ex_para.add_run("Expected Profit Factor:    1.0\n").font.name = 'Courier New'
    ex_para.add_run("Expected Profit = $2,000,000 × (1 - 0.55) × 1.0 = $900,000").font.name = 'Courier New'
    
    doc.add_paragraph()
    
    doc.add_heading("What this means:", level=2)
    doc.add_paragraph(
        "Expected Profit is an estimate based on historical data. Actual results will vary - some treaties "
        "will outperform (lower loss ratios), some will underperform (higher loss ratios). This profit funds:"
    )
    doc.add_paragraph("Operations and staffing costs", style='List Bullet')
    doc.add_paragraph("Claims reserves and contingencies", style='List Bullet')
    doc.add_paragraph("Shareholder returns and capital growth", style='List Bullet')
    doc.add_paragraph("Investment income and returns", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        f"Your portfolio's expected profit is approximately ${portfolio['expected_profit']:,.0f}, "
        f"representing {(portfolio['expected_profit']/portfolio['portfolio_value'])*100:.1f}% of portfolio value. "
        "This is a healthy profit margin."
    )
    
    doc.add_page_break()
    
    add_heading_style(doc, "9. TOTAL LOSS RATIO", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph(
        "The average percentage of premiums that are paid out as claims across your entire portfolio. "
        "It measures claims experience and indicates profitability."
    )
    
    doc.add_heading("Formula:", level=2)
    formula = doc.add_paragraph()
    formula.paragraph_format.left_indent = Inches(0.5)
    formula_run = formula.add_run("Loss Ratio = Total Incurred Losses / Total Premiums")
    formula_run.font.italic = True
    
    doc.add_heading("Interpretation scale:", level=2)
    loss_table = doc.add_table(rows=5, cols=3)
    loss_table.style = 'Light Grid Accent 1'
    
    header_cells = loss_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    shade_cell(header_cells[2], '003d82')
    header_cells[0].text = 'Loss Ratio'
    header_cells[1].text = 'Assessment'
    header_cells[2].text = 'Implication'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    loss_data = [
        ('<50%', '✓✓ Very profitable', 'Claims are well-controlled'),
        ('50-70%', '✓ Healthy/normal', 'Expected reinsurance range'),
        ('70-85%', '⚠️ Challenged', 'High claims affecting profit'),
        ('>85%', '❌ Loss-making', 'Premium not covering claims')
    ]
    
    for i, (ratio, assessment, implication) in enumerate(loss_data, 1):
        row_cells = loss_table.rows[i].cells
        row_cells[0].text = ratio
        row_cells[1].text = assessment
        row_cells[2].text = implication
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
            shade_cell(row_cells[2], 'E8F0F8')
    
    doc.add_paragraph()
    
    doc.add_heading("Your portfolio:", level=2)
    doc.add_paragraph(
        f"Your total loss ratio is approximately {portfolio['total_loss_ratio']:.1%}. This is in the healthy range."
    )
    doc.add_paragraph("Means: For every $1 premium, ~${portfolio['total_loss_ratio']:.2f} goes to claims", style='List Bullet')
    doc.add_paragraph(f"Remaining: ${1-portfolio['total_loss_ratio']:.2f} covers operations, capital, profit", style='List Bullet')
    doc.add_paragraph("This healthy ratio supports your high RORAC", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== LOB BREAKDOWN =====
    add_heading_style(doc, "10. LINE OF BUSINESS (LOB) BREAKDOWN", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph(
        "The distribution of premium across your 5 lines of business. Each LOB has different risk "
        "characteristics, return profiles, and loss patterns."
    )
    
    doc.add_heading("Your portfolio breakdown:", level=2)
    
    lob_table = doc.add_table(rows=6, cols=4)
    lob_table.style = 'Light Grid Accent 1'
    
    header_cells = lob_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    shade_cell(header_cells[2], '003d82')
    shade_cell(header_cells[3], '003d82')
    header_cells[0].text = 'Line of Business'
    header_cells[1].text = 'Premium'
    header_cells[2].text = '% of Portfolio'
    header_cells[3].text = 'Characteristics'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    lob_breakdown = [
        ('Property Catastrophe', 
         f"${portfolio['lob_breakdown'].get('Property Catastrophe', 0)/1e6:.1f}M",
         f"{(portfolio['lob_breakdown'].get('Property Catastrophe', 0)/portfolio['portfolio_value'])*100:.0f}%",
         'High risk, CAT exposure, High returns'),
        ('Casualty',
         f"${portfolio['lob_breakdown'].get('Casualty', 0)/1e6:.1f}M",
         f"{(portfolio['lob_breakdown'].get('Casualty', 0)/portfolio['portfolio_value'])*100:.0f}%",
         'Moderate-high risk, Stable returns'),
        ('Marine & Aviation',
         f"${portfolio['lob_breakdown'].get('Marine & Aviation', 0)/1e6:.1f}M",
         f"{(portfolio['lob_breakdown'].get('Marine & Aviation', 0)/portfolio['portfolio_value'])*100:.0f}%",
         'High risk, CAT exposure, Good returns'),
        ('Financial Lines',
         f"${portfolio['lob_breakdown'].get('Financial Lines', 0)/1e6:.1f}M",
         f"{(portfolio['lob_breakdown'].get('Financial Lines', 0)/portfolio['portfolio_value'])*100:.0f}%",
         'Lower risk, Predictable, Stable returns'),
        ('Specialty',
         f"${portfolio['lob_breakdown'].get('Specialty', 0)/1e6:.1f}M",
         f"{(portfolio['lob_breakdown'].get('Specialty', 0)/portfolio['portfolio_value'])*100:.0f}%",
         'Moderate risk, Niche markets, Competitive')
    ]
    
    for i, (lob, premium, pct, char) in enumerate(lob_breakdown, 1):
        row_cells = lob_table.rows[i].cells
        row_cells[0].text = lob
        row_cells[1].text = premium
        row_cells[2].text = pct
        row_cells[3].text = char
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
            shade_cell(row_cells[2], 'E8F0F8')
            shade_cell(row_cells[3], 'E8F0F8')
    
    doc.add_paragraph()
    
    doc.add_heading("Why this matters:", level=2)
    doc.add_paragraph("Each LOB has different risk/return profiles:")
    doc.add_paragraph("Property Catastrophe: High risk, High potential returns", style='List Bullet')
    doc.add_paragraph("Financial Lines: Lower risk, Lower volatility", style='List Bullet')
    doc.add_paragraph("Your mix balances growth opportunities with portfolio stability", style='List Bullet')
    doc.add_paragraph("Allows specialization by underwriting team members", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== GEOGRAPHIC BREAKDOWN =====
    add_heading_style(doc, "11. GEOGRAPHIC BREAKDOWN", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph(
        "The distribution of premium across 5 geographic regions. Each region has different "
        "economic, regulatory, and natural disaster risk characteristics."
    )
    
    doc.add_heading("Your portfolio breakdown:", level=2)
    
    geo_table = doc.add_table(rows=6, cols=4)
    geo_table.style = 'Light Grid Accent 1'
    
    header_cells = geo_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    shade_cell(header_cells[2], '003d82')
    shade_cell(header_cells[3], '003d82')
    header_cells[0].text = 'Geography'
    header_cells[1].text = 'Premium'
    header_cells[2].text = '% of Portfolio'
    header_cells[3].text = 'Key Characteristics'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    geo_breakdown = [
        ('North America',
         f"${portfolio['geography_breakdown'].get('North America', 0)/1e6:.1f}M",
         f"{(portfolio['geography_breakdown'].get('North America', 0)/portfolio['portfolio_value'])*100:.0f}%",
         'Mature, Stable, Predictable CAT exposure'),
        ('Europe',
         f"${portfolio['geography_breakdown'].get('Europe', 0)/1e6:.1f}M",
         f"{(portfolio['geography_breakdown'].get('Europe', 0)/portfolio['portfolio_value'])*100:.0f}%",
         'Regulated, Sophisticated, Lower but stable'),
        ('Asia Pacific',
         f"${portfolio['geography_breakdown'].get('Asia Pacific', 0)/1e6:.1f}M",
         f"{(portfolio['geography_breakdown'].get('Asia Pacific', 0)/portfolio['portfolio_value'])*100:.0f}%",
         'Emerging, Growth potential, Higher volatility'),
        ('Latin America',
         f"${portfolio['geography_breakdown'].get('Latin America', 0)/1e6:.1f}M",
         f"{(portfolio['geography_breakdown'].get('Latin America', 0)/portfolio['portfolio_value'])*100:.0f}%",
         'Growth opportunity, Higher volatility'),
        ('Africa/Middle East',
         f"${portfolio['geography_breakdown'].get('Africa/Middle East', 0)/1e6:.1f}M",
         f"{(portfolio['geography_breakdown'].get('Africa/Middle East', 0)/portfolio['portfolio_value'])*100:.0f}%",
         'Small exposure, Diversification benefit')
    ]
    
    for i, (geo, premium, pct, char) in enumerate(geo_breakdown, 1):
        row_cells = geo_table.rows[i].cells
        row_cells[0].text = geo
        row_cells[1].text = premium
        row_cells[2].text = pct
        row_cells[3].text = char
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
            shade_cell(row_cells[2], 'E8F0F8')
            shade_cell(row_cells[3], 'E8F0F8')
    
    doc.add_paragraph()
    
    doc.add_heading("Why this matters:", level=2)
    doc.add_paragraph("Different regions have different risk profiles:")
    doc.add_paragraph("Natural disasters (hurricanes, earthquakes) are regional", style='List Bullet')
    doc.add_paragraph("Political/economic risks vary significantly", style='List Bullet')
    doc.add_paragraph("Your $64M spread across 5 regions reduces geographic concentration", style='List Bullet')
    doc.add_paragraph("Provides growth opportunities while managing systemic risk", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== TREATIES DETAILS =====
    add_heading_style(doc, "12. TREATIES DETAILS (50 Individual Contracts)", 1)
    
    doc.add_heading("What each column means:", level=2)
    
    treaties_table = doc.add_table(rows=15, cols=2)
    treaties_table.style = 'Light Grid Accent 1'
    
    header_cells = treaties_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    header_cells[0].text = 'Column'
    header_cells[1].text = 'What It Means'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    treaty_cols = [
        ('Treaty ID', 'Unique identifier for each contract (e.g., TR-Pro-0001)'),
        ('LOB', 'Line of Business type'),
        ('Geography', 'Region where risk is located'),
        ('Treaty Type', 'How risk is shared (Quota Share, XL, etc.)'),
        ('Premium', 'Total premium for this contract'),
        ('Incurred Loss', 'Expected or actual claims amount'),
        ('Loss Ratio', 'Incurred Loss / Premium'),
        ('Capital Required', 'Risk capital allocated to this treaty'),
        ('Expected Profit', 'Estimated profit from this treaty'),
        ('RORAC', 'Return On Risk-Adjusted Capital (%)'),
        ('Ceded Premium', 'Premium passed to other reinsurers'),
        ('Renewal Date', 'When contract expires and renews'),
        ('Performance', 'Rating (Excellent, Good, Adequate, Watch, Poor)'),
        ('Rating', 'Credit rating of counterparty')
    ]
    
    for i, (col, meaning) in enumerate(treaty_cols, 1):
        row_cells = treaties_table.rows[i].cells
        row_cells[0].text = col
        row_cells[1].text = meaning
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
    
    doc.add_paragraph()
    
    doc.add_heading("Why analyze individual treaties:", level=2)
    doc.add_paragraph("Individual treaty analysis allows you to:")
    doc.add_paragraph("Identify underperforming or problem contracts", style='List Bullet')
    doc.add_paragraph("Spot renewal opportunities and renegotiation potential", style='List Bullet')
    doc.add_paragraph("Monitor claims development over time", style='List Bullet')
    doc.add_paragraph("Track profitability per contract and per underwriter", style='List Bullet')
    doc.add_paragraph("Make informed decisions about portfolio optimization", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== PERFORMANCE DISTRIBUTION =====
    add_heading_style(doc, "13. PERFORMANCE DISTRIBUTION", 1)
    
    doc.add_heading("What it is:", level=2)
    doc.add_paragraph(
        "A count of how many treaties fall into each performance rating category. Shows overall "
        "portfolio health from a performance perspective."
    )
    
    doc.add_heading("Performance ratings explained:", level=2)
    
    perf_table = doc.add_table(rows=6, cols=4)
    perf_table.style = 'Light Grid Accent 1'
    
    header_cells = perf_table.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    shade_cell(header_cells[2], '003d82')
    shade_cell(header_cells[3], '003d82')
    header_cells[0].text = 'Status'
    header_cells[1].text = 'Definition'
    header_cells[2].text = 'Action'
    header_cells[3].text = 'Expected % of Portfolio'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    perf_data = [
        ('Excellent', 'Outperforming expectations, low claims', 'Renew, expand', '10-20%'),
        ('Good', 'Performing as expected, stable results', 'Maintain, grow', '30-40%'),
        ('Adequate', 'Meeting minimum standards', 'Monitor, stable', '20-35%'),
        ('Watch', 'Some concern, elevated claims', 'Close monitoring, negotiate', '10-15%'),
        ('Poor', 'Significant underperformance', 'Exit or restructure', '<5%')
    ]
    
    for i, (status, definition, action, expected) in enumerate(perf_data, 1):
        row_cells = perf_table.rows[i].cells
        row_cells[0].text = status
        row_cells[1].text = definition
        row_cells[2].text = action
        row_cells[3].text = expected
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
            shade_cell(row_cells[2], 'E8F0F8')
            shade_cell(row_cells[3], 'E8F0F8')
    
    doc.add_paragraph()
    
    doc.add_heading("Healthy portfolio structure:", level=2)
    if portfolio['performance_distribution']:
        perf_dist = portfolio['performance_distribution']
        total_treaties = sum(perf_dist.values())
        
        for status, count in perf_dist.items():
            pct = (count / total_treaties) * 100
            doc.add_paragraph(f"{status}: {count} treaties ({pct:.0f}%)", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph("This distribution shows:")
    doc.add_paragraph("✓ Core performing treaties (Excellent + Good) drive profits", style='List Bullet')
    doc.add_paragraph("✓ Adequate treaties provide stable baseline", style='List Bullet')
    doc.add_paragraph("✓ Watch/Poor are minority requiring active management", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== KEY TAKEAWAYS =====
    add_heading_style(doc, "14. KEY TAKEAWAYS & ASSESSMENT", 1)
    
    doc.add_heading("Portfolio Strengths ✓", level=2)
    doc.add_paragraph("Exceptional RORAC (229%) - very profitable", style='List Bullet')
    doc.add_paragraph("Good diversification (0.75) across LOBs", style='List Bullet')
    doc.add_paragraph("Balanced geographic spread", style='List Bullet')
    doc.add_paragraph("50 treaties provide excellent risk distribution", style='List Bullet')
    doc.add_paragraph("Acceptable capital utilization (102%)", style='List Bullet')
    doc.add_paragraph("Healthy expected profit margins", style='List Bullet')
    doc.add_paragraph("Controlled loss ratios (50-65% range)", style='List Bullet')
    
    doc.add_heading("Areas to Monitor ⚠️", level=2)
    doc.add_paragraph("Slight over-leverage (102% > 100%) - may need more capital", style='List Bullet')
    doc.add_paragraph("Property Catastrophe concentration at 30% - monitor CAT risk", style='List Bullet')
    doc.add_paragraph("North America concentration at 40% - monitor regional risks", style='List Bullet')
    doc.add_paragraph("Some Watch/Poor performing treaties - require active management", style='List Bullet')
    
    doc.add_heading("Portfolio Direction 📈", level=2)
    doc.add_paragraph("This is a HEALTHY, PROFITABLE portfolio")
    doc.add_paragraph("Ready for growth or optimization")
    doc.add_paragraph("Well-positioned for scenario analysis")
    doc.add_paragraph("Can support strategic expansion initiatives")
    
    doc.add_heading("Recommended Next Steps:", level=2)
    doc.add_paragraph("Analyze individual treaty performance for renewal decisions", style='List Bullet')
    doc.add_paragraph("Consider capital raise to support 102% utilization", style='List Bullet')
    doc.add_paragraph("Monitor Property Catastrophe exposure quarterly", style='List Bullet')
    doc.add_paragraph("Review Watch/Poor performing contracts for restructuring", style='List Bullet')
    doc.add_paragraph("Model growth scenarios using different rate/volume assumptions", style='List Bullet')
    doc.add_paragraph("Stress test portfolio for major CAT events", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ===== APPENDIX: QUICK REFERENCE =====
    add_heading_style(doc, "APPENDIX: QUICK REFERENCE GUIDE", 1)
    
    doc.add_heading("Key Metrics at a Glance:", level=2)
    
    quick_ref = doc.add_table(rows=10, cols=2)
    quick_ref.style = 'Light Grid Accent 1'
    
    header_cells = quick_ref.rows[0].cells
    shade_cell(header_cells[0], '003d82')
    shade_cell(header_cells[1], '003d82')
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    quick_data = [
        ('Portfolio Value', f"${portfolio['portfolio_value']:,.0f}"),
        ('Capital Utilization', f"{portfolio['capital_utilization']:.1f}%"),
        ('Average RORAC', f"{portfolio['avg_rorac']:.1f}%"),
        ('Diversification Score', f"{portfolio['diversification_score']:.2f}"),
        ('Total Loss Ratio', f"{portfolio['total_loss_ratio']:.1%}"),
        ('Expected Profit', f"${portfolio['expected_profit']:,.0f}"),
        ('Total Capital Required', f"${portfolio['total_capital']:,.0f}"),
        ('Treaty Count', f"{portfolio['treaty_count']}"),
        ('Document Created', datetime.now().strftime('%Y-%m-%d %H:%M'))
    ]
    
    for i, (metric, value) in enumerate(quick_data, 1):
        row_cells = quick_ref.rows[i].cells
        row_cells[0].text = metric
        row_cells[1].text = value
        if i % 2 == 0:
            shade_cell(row_cells[0], 'E8F0F8')
            shade_cell(row_cells[1], 'E8F0F8')
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'Portfolio_Summary_Reference.docx')
    doc.save(output_path)
    
    print(f"\n✅ Portfolio reference document created successfully!")
    print(f"📁 File saved to: {output_path}")
    print(f"\n📄 Document contains:")
    print(f"   • Executive Summary")
    print(f"   • Detailed explanations of all 9 key metrics")
    print(f"   • Industry context and benchmarks")
    print(f"   • LOB and Geographic analysis")
    print(f"   • Treaties details guide")
    print(f"   • Performance distribution analysis")
    print(f"   • Key takeaways and recommendations")
    print(f"   • Quick reference guide")
    print(f"\n📊 Total pages: ~25 pages")
    print(f"🎨 Professional formatting with color-coded headers and tables")
    
    return output_path

if __name__ == '__main__':
    create_portfolio_reference_document()
