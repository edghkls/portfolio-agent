"""
Generate Portfolio Terms Reference as Word Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Title
title = doc.add_heading('Portfolio Details - Complete Reference Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(0, 51, 102)

# Subtitle
subtitle = doc.add_paragraph('Reinsurance Portfolio Optimization Dashboard')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(12)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Section 1: Local Access
doc.add_heading('1. Local Access', level=1)
p = doc.add_paragraph()
p.add_run('URL: ').bold = True
p.add_run('http://localhost:5001/portfolio')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Files:\n').bold = True
p.add_run('• Frontend: web_ui/templates/portfolio-dynamic.html\n')
p.add_run('• Backend: web_ui/app_simple.py\n')
p.add_run('• Data Source: data_connectors/mock_portfolio.py')

doc.add_page_break()

# Section 2: Core Treaty Information
doc.add_heading('2. Core Treaty Information', level=1)

# Treaty ID
doc.add_heading('Treaty ID', level=2)
doc.add_paragraph('Unique identifier/number assigned to each reinsurance contract for tracking and management purposes.')
p = doc.add_paragraph('Example: TR-2026-0042')
p.paragraph_format.left_indent = Inches(0.25)

# LOB
doc.add_heading('LOB (Line of Business)', level=2)
doc.add_paragraph('Category of insurance coverage that defines the type of risk being insured.')
p = doc.add_paragraph('Lines of Business:')
p.paragraph_format.left_indent = Inches(0.25)
for lob in ['Property Catastrophe', 'Casualty', 'Marine & Aviation', 'Financial Lines', 'Specialty']:
    p = doc.add_paragraph(lob, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

# Geography
doc.add_heading('Geography', level=2)
doc.add_paragraph('Geographic region where the insurance risk is located - determines regulatory environment and risk profile.')
p = doc.add_paragraph('Regions:')
p.paragraph_format.left_indent = Inches(0.25)
for geo in ['North America', 'Europe', 'Asia Pacific', 'Latin America', 'Africa/Middle East']:
    p = doc.add_paragraph(geo, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# Section 3: Financial Metrics
doc.add_heading('3. Financial Metrics', level=1)

# Premium
doc.add_heading('Premium', level=2)
doc.add_paragraph('Money received from clients for providing insurance coverage over the contract period.')
doc.add_paragraph('Example: $8.5M received for one-year coverage')
doc.add_paragraph('• Higher premiums = more revenue from clients')
doc.add_paragraph('• Represents the top-line income before any claims')

# RORAC
doc.add_heading('RORAC (Return On Risk-Adjusted Capital)', level=2)
doc.add_paragraph('Percentage return earned per dollar of capital invested - measures how efficiently capital is being used.')
doc.add_paragraph('Formula: (Profit / Risk-Adjusted Capital) × 100')
doc.add_paragraph('Example: 22.5% means earning $0.225 for every $1 of capital allocated')
doc.add_paragraph('• Good RORAC: 20%+')
doc.add_paragraph('• Poor RORAC: Below 10%')
doc.add_paragraph('• Used to compare profitability between different treaties')

# Loss Ratio
doc.add_heading('Loss Ratio', level=2)
doc.add_paragraph('Percentage of premiums paid out as insurance claims - indicates claim severity and underwriting accuracy.')
doc.add_paragraph('Formula: (Claims Paid / Premium Received) × 100')
doc.add_paragraph('Example: 45% loss ratio on $8.5M premium = $3.8M in claims paid')
doc.add_paragraph('• Excellent: Below 50% (profitable)')
doc.add_paragraph('• Acceptable: 50-70% (still profitable but higher risk)')
doc.add_paragraph('• Warning: 70%+ (paying too much in claims, need to review)')

# Expected Profit
doc.add_heading('Expected Profit', level=2)
doc.add_paragraph('Forecasted earnings after all claims and operational costs - bottom-line profitability.')
doc.add_paragraph('Formula: Premium - Expected Claims - Operating Costs')
doc.add_paragraph('Example: $8.5M (Premium) - $3.8M (Claims) - $0.5M (Costs) = $4.2M Profit')
doc.add_paragraph('• Positive profit means the contract is financially viable')
doc.add_paragraph('• Used to assess contract attractiveness')

doc.add_page_break()

# Section 4: Risk & Quality Metrics
doc.add_heading('4. Risk & Quality Metrics', level=1)

# Rating
doc.add_heading('Rating', level=2)
doc.add_paragraph('Credit/financial strength rating indicating the likelihood of claims being paid without default.')
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Rating'
header_cells[1].text = 'Meaning'
shade_cell(header_cells[0], '0033CC')
shade_cell(header_cells[1], '0033CC')
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

ratings = [
    ('AAA', 'Exceptional - virtually no risk'),
    ('AA', 'Very Strong - very low risk'),
    ('A', 'Strong - low risk'),
    ('BBB', 'Adequate - moderate risk'),
    ('BB', 'Speculative - higher risk'),
    ('C', 'Highly Risky - potential default risk')
]

for i, (rating, meaning) in enumerate(ratings, 1):
    row = table.rows[i].cells
    row[0].text = rating
    row[1].text = meaning

# Underwriter
doc.add_heading('Underwriter', level=2)
doc.add_paragraph('Person or team responsible for assessing, approving, and managing the insurance contract.')
doc.add_paragraph('• Evaluates risk')
doc.add_paragraph('• Sets premium rates')
doc.add_paragraph('• Monitors contract performance')
doc.add_paragraph('Example: John Smith, Mary Johnson')

# Performance Status
doc.add_heading('Performance Status', level=2)
doc.add_paragraph('Current health/condition of the contract indicating how well it\'s performing against expectations.')
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
headers = ['Status', 'Meaning', 'Action']
for i, header in enumerate(headers):
    header_cells[i].text = header
    shade_cell(header_cells[i], '0033CC')
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

statuses = [
    ('Adequate', 'Contract performing as expected - no issues', 'Monitor regularly'),
    ('Warning', 'Performance declining or risk emerging - needs attention', 'Review and adjust'),
    ('Critical', 'Serious problems - contract at risk - urgent action needed', 'Immediate intervention')
]

for i, (status, meaning, action) in enumerate(statuses, 1):
    row = table.rows[i].cells
    row[0].text = status
    row[1].text = meaning
    row[2].text = action

doc.add_page_break()

# Section 5: Data Categories - LOB
doc.add_heading('5. Data Categories - Line of Business (LOB)', level=1)

lobs_detail = [
    {
        'name': '1. Property Catastrophe',
        'desc': 'Insurance for major natural disaster coverage - protects against large-scale catastrophic events.',
        'covers': ['Hurricanes', 'Earthquakes', 'Floods', 'Tornados', 'Wildfires'],
        'premium': 'High (due to high potential losses)',
        'loss_ratio': 'Variable (depends on natural disaster occurrence)',
        'key': 'Return period of events (1-in-100 year, 1-in-200 year)',
        'example': 'Treaty covers property damage in hurricane-prone areas'
    },
    {
        'name': '2. Casualty',
        'desc': 'Insurance for accidents, injuries, and legal liability - covers bodily injury and property damage claims.',
        'covers': ['Auto accidents', 'Workplace injuries', 'Slip & fall', 'Legal liability'],
        'premium': 'Moderate',
        'loss_ratio': 'More predictable (historical data available)',
        'key': 'Claims per policy, average claim cost',
        'example': 'Liability coverage for commercial businesses'
    },
    {
        'name': '3. Marine & Aviation',
        'desc': 'Insurance for ships, aircraft, cargo, and maritime risks - specialized insurance for transportation.',
        'covers': ['Cargo loss/damage', 'Ship collisions', 'Aircraft accidents', 'Piracy'],
        'premium': 'High (specialized risks)',
        'loss_ratio': 'Volatile (rare but severe events)',
        'key': 'Shipping routes, aircraft types',
        'example': 'Coverage for international cargo shipments'
    },
    {
        'name': '4. Financial Lines',
        'desc': 'Insurance for financial losses and professional liability - protects financial institutions and professionals.',
        'covers': ['Directors & Officers liability', 'Professional liability', 'Financial crime', 'Errors & omissions'],
        'premium': 'Moderate to High',
        'loss_ratio': 'Lower than property (fewer but larger claims)',
        'key': 'Industry type, claim history',
        'example': 'Coverage for bank executives against litigation risk'
    },
    {
        'name': '5. Specialty',
        'desc': 'Niche/specialized insurance for unique risks - covers non-standard risks requiring expertise.',
        'covers': ['Cyber attacks', 'Kidnap & ransom', 'Political risk', 'Product recall', 'Event cancellation'],
        'premium': 'Highly variable',
        'loss_ratio': 'Unpredictable (limited historical data)',
        'key': 'Risk-specific factors',
        'example': 'Coverage for data breach liability'
    }
]

for lob in lobs_detail:
    doc.add_heading(lob['name'], level=2)
    doc.add_paragraph(lob['desc'])
    
    p = doc.add_paragraph()
    p.add_run('Covers: ').bold = True
    p.add_run(', '.join(lob['covers']))
    
    p = doc.add_paragraph()
    p.add_run('Premium Range: ').bold = True
    p.add_run(lob['premium'])
    
    p = doc.add_paragraph()
    p.add_run('Loss Ratio: ').bold = True
    p.add_run(lob['loss_ratio'])
    
    p = doc.add_paragraph()
    p.add_run('Key Metric: ').bold = True
    p.add_run(lob['key'])
    
    p = doc.add_paragraph()
    p.add_run('Example: ').bold = True
    p.add_run(lob['example'])
    
    doc.add_paragraph()

doc.add_page_break()

# Section 6: Geographic Regions
doc.add_heading('6. Geographic Regions', level=1)

geos = [
    {
        'name': 'North America',
        'desc': 'USA, Canada, Mexico - Most developed insurance market',
        'char': 'Large premiums, strict regulation, lower loss ratios',
        'risk': 'Well-established, predictable',
        'market': 'Most competitive, highest volume',
        'example': 'Auto, property, workers compensation'
    },
    {
        'name': 'Europe',
        'desc': 'EU countries, UK, Switzerland - Highly regulated insurance market',
        'char': 'Moderate premiums, strong regulatory oversight, organized market',
        'risk': 'Stable, compliant',
        'market': 'Mature, consolidating',
        'example': 'Motor, liability, health'
    },
    {
        'name': 'Asia Pacific',
        'desc': 'Japan, Australia, Singapore, India, China - Growing rapidly',
        'char': 'High premiums, emerging risks, regulatory variation',
        'risk': 'Growing but less predictable',
        'market': 'Fastest growing, high potential',
        'example': 'Natural disaster, cyber, auto'
    },
    {
        'name': 'Latin America',
        'desc': 'Brazil, Argentina, Chile, Colombia, Mexico - Emerging market',
        'char': 'Moderate to high premiums, developing infrastructure, variable regulation',
        'risk': 'Moderate risk, improving',
        'market': 'Growth opportunity',
        'example': 'Property, casualty, specialty'
    },
    {
        'name': 'Africa/Middle East',
        'desc': 'African continent, Saudi Arabia, UAE, Israel - Specialized market',
        'char': 'Highest premiums, specialized underwriting needed, limited competition',
        'risk': 'Higher risk, requires expertise',
        'market': 'Niche, specialized knowledge required',
        'example': 'Political risk, property, specialty'
    }
]

for geo in geos:
    doc.add_heading(geo['name'], level=2)
    doc.add_paragraph(geo['desc'])
    
    p = doc.add_paragraph()
    p.add_run('Characteristics: ').bold = True
    p.add_run(geo['char'])
    
    p = doc.add_paragraph()
    p.add_run('Risk Profile: ').bold = True
    p.add_run(geo['risk'])
    
    p = doc.add_paragraph()
    p.add_run('Market: ').bold = True
    p.add_run(geo['market'])
    
    p = doc.add_paragraph()
    p.add_run('Example Treaties: ').bold = True
    p.add_run(geo['example'])
    
    doc.add_paragraph()

doc.add_page_break()

# Section 7: Example Treaty
doc.add_heading('7. Example Treaty Breakdown', level=1)
doc.add_heading('Treaty Details', level=2)

example_data = [
    ('Treaty ID', 'TR-2026-0042'),
    ('Line of Business', 'Property Catastrophe'),
    ('Geography', 'Asia Pacific'),
    ('Premium', '$8.5M'),
    ('RORAC', '22.5%'),
    ('Loss Ratio', '45%'),
    ('Expected Profit', '$4.2M'),
    ('Rating', 'AA'),
    ('Underwriter', 'John Smith'),
    ('Performance Status', 'Adequate')
]

table = doc.add_table(rows=len(example_data) + 1, cols=2)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Field'
header_cells[1].text = 'Value'
shade_cell(header_cells[0], '0033CC')
shade_cell(header_cells[1], '0033CC')
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (field, value) in enumerate(example_data, 1):
    row = table.rows[i].cells
    row[0].text = field
    row[1].text = value

doc.add_heading('What This Means', level=2)

interpretation = [
    ('Treaty ID', 'TR-2026-0042', 'This is contract #42 from 2026'),
    ('LOB', 'Property Catastrophe', 'Covers major disasters (hurricanes, earthquakes)'),
    ('Geography', 'Asia Pacific', 'Serves customers in Japan, Australia, Singapore, India'),
    ('Premium', '$8.5M', 'Received $8.5M from clients for this coverage'),
    ('RORAC', '22.5%', 'Earning 22.5% return on invested capital = EXCELLENT'),
    ('Loss Ratio', '45%', 'Paying out 45% of premium as claims = GOOD (profitable)'),
    ('Expected Profit', '$4.2M', 'Forecast profit = $4.2M (Premium - Claims - Costs)'),
    ('Rating', 'AA', 'Very Strong financial rating = LOW DEFAULT RISK'),
    ('Underwriter', 'John Smith', 'John is responsible for managing this contract'),
    ('Status', 'Adequate', 'Contract is performing as expected = NO ISSUES')
]

table = doc.add_table(rows=len(interpretation) + 1, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
headers = ['Field', 'Value', 'Interpretation']
for i, header in enumerate(headers):
    header_cells[i].text = header
    shade_cell(header_cells[i], '0033CC')
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (field, value, interp) in enumerate(interpretation, 1):
    row = table.rows[i].cells
    row[0].text = field
    row[1].text = value
    row[2].text = interp

doc.add_page_break()

# Section 8: Why Metrics Matter
doc.add_heading('8. Why These Metrics Matter', level=1)

doc.add_heading('RORAC (Return On Risk-Adjusted Capital)', level=2)
doc.add_paragraph('Why: Measures how efficiently capital is being deployed')
doc.add_paragraph('Good: 20%+ = capital is working hard, generating strong returns')
doc.add_paragraph('Poor: Below 10% = capital underutilized, consider reallocating')
doc.add_paragraph('Action: Move capital from low RORAC to high RORAC treaties')

doc.add_heading('Loss Ratio', level=2)
doc.add_paragraph('Why: Indicates underwriting quality and claim experience')
doc.add_paragraph('Excellent: Below 50% = underwriting is accurate, clients are good risks')
doc.add_paragraph('Warning: 70%+ = either prices are too low OR clients are riskier than expected')
doc.add_paragraph('Action: Increase premiums or tighten underwriting standards for high-ratio contracts')

doc.add_heading('Rating (Credit Quality)', level=2)
doc.add_paragraph('Why: Predicts likelihood of being able to pay claims when due')
doc.add_paragraph('Strong (AA): Unlikely to default, safe')
doc.add_paragraph('Weak (C): May default on claims, risky')
doc.add_paragraph('Action: Avoid concentrating portfolio in low-rated counterparties')

doc.add_heading('Performance Status', level=2)
doc.add_paragraph('Why: Early warning indicator of contract health')
doc.add_paragraph('Adequate: No intervention needed, monitor normally')
doc.add_paragraph('Warning: Problems emerging, needs review and adjustment')
doc.add_paragraph('Critical: Urgent action required to prevent major loss')
doc.add_paragraph('Action: Critical treaties require immediate management attention')

doc.add_page_break()

# Section 9: Healthy Portfolio Ranges
doc.add_heading('9. Quick Reference - Healthy Portfolio Ranges', level=1)

ranges = [
    ('RORAC', '20%+', '15-20%', 'Below 15%'),
    ('Loss Ratio', '30-50%', '50-70%', 'Above 70%'),
    ('Rating', 'AA-AAA', 'A-BBB', 'BB-C'),
    ('% Adequate Status', '85%+', '70-85%', 'Below 70%'),
    ('Profit Margin', '40%+', '25-40%', 'Below 25%')
]

table = doc.add_table(rows=len(ranges) + 1, cols=4)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
headers = ['Metric', 'Excellent', 'Acceptable', 'Warning']
for i, header in enumerate(headers):
    header_cells[i].text = header
    shade_cell(header_cells[i], '0033CC')
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (metric, excellent, acceptable, warning) in enumerate(ranges, 1):
    row = table.rows[i].cells
    row[0].text = metric
    row[1].text = excellent
    row[2].text = acceptable
    row[3].text = warning

# Footer
doc.add_page_break()
footer = doc.add_paragraph()
footer.add_run('Document Information\n').bold = True
footer.add_run('Last Updated: July 15, 2026\n')
footer.add_run('Portfolio Agent Version: 1.0\n')
footer.add_run('Status: Production Ready\n')
footer.add_run('\nFor technical support, contact: Portfolio Agent Development Team')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
doc.save('Portfolio_Terms_Reference.docx')
print("✅ Word document created successfully: Portfolio_Terms_Reference.docx")
