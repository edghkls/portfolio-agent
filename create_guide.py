from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create Document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Title
title = doc.add_heading('Dynamic Portfolio Optimization Dashboard', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_format = title.runs[0]
title_format.font.color.rgb = RGBColor(0, 61, 130)  # Primary color

# Subtitle
subtitle = doc.add_paragraph('Complete Guide to Portfolio Management Metrics & Terminology')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(12)
subtitle_format.font.italic = True

# Date
date_para = doc.add_paragraph(f'Document Date: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.size = Pt(10)
date_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc_items = [
    'What is GWP and Why It\'s Used',
    'Core Performance Metrics',
    'Strategic Performance Metrics',
    'Risk Assessment Metrics',
    'Geographic & LOB Analysis',
    'Capital Efficiency & ROI',
    'Risk-Return Optimization',
    'Dashboard Benefits Summary'
]
for i, item in enumerate(toc_items, 1):
    doc.add_paragraph(f'{i}. {item}', style='List Bullet')

doc.add_page_break()

# 1. GWP Section
doc.add_heading('1. What is GWP and Why It\'s Used?', 1)
doc.add_heading('Definition:', 2)
p = doc.add_paragraph('GWP stands for ')
p.add_run('Gross Written Premium').bold = True
p.add_run(' - the total amount of insurance premiums written by the company in a specific period (monthly, quarterly, or annually), ')
p.add_run('before any reinsurance or claims payouts').italic = True
p.add_run('.')

doc.add_heading('Simple Explanation:', 2)
doc.add_paragraph(
    'Think of GWP as the total "sales revenue" of the insurance company. If you sell 100 insurance policies with an average premium of $2.5M per policy, your GWP = $250M. This is the money collected from customers.',
    style='List Bullet'
)

doc.add_heading('Why GWP is Used in This Dashboard:', 2)
reasons = [
    'Revenue Indicator: Shows the business volume and market reach',
    'Growth Metric: Tracks if the company is growing (12% growth = gaining more customers)',
    'Risk Exposure: Higher GWP = more risk exposure that needs capital reserves',
    'Performance Baseline: Used to calculate other metrics like loss ratio and RORAC',
    'Market Position: Indicates competitive standing and market share'
]
for reason in reasons:
    doc.add_paragraph(reason, style='List Number')

doc.add_heading('Current Dashboard Data:', 2)
doc.add_paragraph('GWP: $256.9M with ↑ +5.2% growth', style='List Bullet')
p = doc.add_paragraph()
p.add_run('Meaning: ').bold = True
p.add_run('The company collected $256.9M in premiums this period, up 5.2% from last period. This is positive growth!')

doc.add_page_break()

# 2. Core Performance Metrics
doc.add_heading('2. Core Performance Metrics', 1)

metrics_data = [
    {
        'name': '💰 GWP (Gross Written Premium)',
        'value': '$256.9M ↑ +5.2%',
        'simple': 'Total premiums collected from all insurance policies',
        'why': 'Measures revenue and business volume',
        'good_when': 'Increasing (positive growth)'
    },
    {
        'name': '🎯 Capital Utilization (Capital Util)',
        'value': '98.9%',
        'simple': 'Percentage of available capital being actively deployed in insurance business',
        'why': 'Shows how efficiently capital is being used. 100% = all capital is working',
        'good_when': 'Between 75-95% (98.9% means we\'re fully utilized, which is excellent but tight)'
    },
    {
        'name': '📈 RORAC (Return on Risk-Adjusted Capital)',
        'value': '23.1%',
        'simple': 'How much profit we make for every $1 of risk capital deployed',
        'why': 'Shows profitability relative to risk. Higher = better returns for the risk taken',
        'good_when': 'Above 20% (23.1% is excellent!)'
    },
    {
        'name': '🔗 Diversification Score',
        'value': '0.76 / 1.0',
        'simple': 'Measures how spread out our portfolio is (1.0 = perfectly diversified)',
        'why': 'Lower concentration = lower risk. 0.76 means good diversification',
        'good_when': 'Above 0.70 (0.76 is very good)'
    },
    {
        'name': '🚀 Premium Growth',
        'value': '12.0% ↑ +8.3%',
        'simple': 'Year-over-year increase in total premiums written',
        'why': 'Shows business expansion and market success',
        'good_when': 'Double digits (12% is excellent growth)'
    }
]

for metric in metrics_data:
    doc.add_heading(metric['name'], 2)
    
    p = doc.add_paragraph()
    p.add_run('Current Value: ').bold = True
    p.add_run(metric['value'])
    
    p = doc.add_paragraph()
    p.add_run('In Simple Terms: ').bold = True
    p.add_run(metric['simple'])
    
    p = doc.add_paragraph()
    p.add_run('Why It Matters: ').bold = True
    p.add_run(metric['why'])
    
    p = doc.add_paragraph()
    p.add_run('Good When: ').bold = True
    p.add_run(metric['good_when'])
    
    doc.add_paragraph()

doc.add_page_break()

# 3. Strategic Performance Metrics
doc.add_heading('3. Strategic Performance Metrics', 1)

strategic_metrics = [
    {
        'name': '⚠️ Claims Ratio',
        'value': '58.0%',
        'simple': 'Percentage of premiums paid out as claims. 58% = for every $100 collected, $58 paid in claims',
        'why': 'Shows underwriting quality. Lower = better (means fewer/smaller claims)',
        'good_when': 'Below 60% (58% is excellent - leaves room for profit)'
    },
    {
        'name': '⚡ Capital Efficiency',
        'value': '$0.85K per $1 of capital',
        'simple': 'How much profit each dollar of capital generates',
        'why': 'Measures effective capital deployment',
        'good_when': 'Above $0.75 (higher = better capital efficiency)'
    },
    {
        'name': '💼 UW Performance (Underwriting)',
        'value': '85% ↑ +3.1%',
        'simple': 'Quality score for how well underwriters are selecting and pricing policies',
        'why': 'Shows quality of risk selection. Higher = better judgement calls',
        'good_when': 'Above 80% (85% means excellent underwriting decisions)'
    },
    {
        'name': '📊 Risk Score',
        'value': '100/100',
        'simple': 'Overall assessment of portfolio risk level (100 = maximum capacity)',
        'why': 'Comprehensive risk measure showing solvency and capacity',
        'good_when': 'At or near 100 (shows full capacity and strength)'
    },
    {
        'name': '💡 Efficiency Score',
        'value': '33/100',
        'simple': 'Operational efficiency in processing, claims, and administration',
        'why': 'Lower = more efficient (room for optimization)',
        'good_when': 'Growing upward (33 is baseline with room for improvement to 40+)'
    }
]

for metric in strategic_metrics:
    doc.add_heading(metric['name'], 2)
    
    p = doc.add_paragraph()
    p.add_run('Current Value: ').bold = True
    p.add_run(metric['value'])
    
    p = doc.add_paragraph()
    p.add_run('In Simple Terms: ').bold = True
    p.add_run(metric['simple'])
    
    p = doc.add_paragraph()
    p.add_run('Why It Matters: ').bold = True
    p.add_run(metric['why'])
    
    p = doc.add_paragraph()
    p.add_run('Good When: ').bold = True
    p.add_run(metric['good_when'])
    
    doc.add_paragraph()

doc.add_page_break()

# 4. Risk Assessment
doc.add_heading('4. Risk Assessment Metrics', 1)

doc.add_heading('Portfolio Health Status:', 2)
risks = [
    ('Concentration Risk', 'Low', 'Diversification score 0.76 - portfolio well spread'),
    ('Catastrophe Exposure', 'Low', 'Property CAT exposure at historical average'),
    ('Capital Utilization', 'Low', 'Capital deployment within target range (98.9%)')
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Risk Type'
hdr_cells[1].text = 'Level'
hdr_cells[2].text = 'Detail'

for risk_name, level, detail in risks:
    row_cells = table.add_row().cells
    row_cells[0].text = risk_name
    row_cells[1].text = level
    row_cells[2].text = detail

doc.add_paragraph()

doc.add_heading('Key Risk Indicators:', 2)
doc.add_paragraph('Capital: ⚠️ Adequate (98.9% utilization)', style='List Bullet')
doc.add_paragraph('Concentration: ✅ Low (well diversified)', style='List Bullet')
doc.add_paragraph('Claims: ✅ Excellent (58% ratio)', style='List Bullet')
doc.add_paragraph('Market: ✅ Good (growing 12%)', style='List Bullet')

doc.add_page_break()

# 5. Geographic & LOB Analysis
doc.add_heading('5. Geographic & Line of Business (LOB) Analysis', 1)

doc.add_heading('Lines of Business Breakdown:', 2)
lobs = [
    ('Property', '$45.2M', '17.5%', 'Building/structure insurance'),
    ('Casualty', '$85.1M', '33.0%', 'Liability and injury claims - LARGEST segment'),
    ('Marine', '$32.1M', '12.4%', 'Shipping and maritime'),
    ('Specialty', '$54.2M', '21.0%', 'High-value specialized risks'),
    ('Reinsurance', '$40.2M', '15.5%', 'Insurance for other insurers')
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'LOB'
hdr_cells[1].text = 'Premium'
hdr_cells[2].text = '% Share'
hdr_cells[3].text = 'Description'

for lob, premium, share, desc in lobs:
    row_cells = table.add_row().cells
    row_cells[0].text = lob
    row_cells[1].text = premium
    row_cells[2].text = share
    row_cells[3].text = desc

doc.add_paragraph()

doc.add_heading('Geographic Distribution:', 2)
geos = [
    ('North America', '$102.0M', '39.6%', '52M profit - PRIMARY MARKET'),
    ('Europe', '$64.0M', '24.8%', '38M profit'),
    ('Asia Pacific', '$51.0M', '19.7%', '28M profit - FASTEST GROWING'),
    ('Latin America', '$26.0M', '10.1%', '12M profit - EMERGING'),
    ('Africa/ME', '$13.0M', '5.1%', '8M profit - EXPANSION POTENTIAL')
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Region'
hdr_cells[1].text = 'Premium'
hdr_cells[2].text = '% Share'
hdr_cells[3].text = 'Notes'

for geo, prem, share, note in geos:
    row_cells = table.add_row().cells
    row_cells[0].text = geo
    row_cells[1].text = prem
    row_cells[2].text = share
    row_cells[3].text = note

doc.add_page_break()

# 6. Capital Efficiency & ROI
doc.add_heading('6. Capital Efficiency & ROI Deep Dive', 1)

doc.add_heading('Capital Allocation Breakdown:', 2)
doc.add_paragraph('Total Available Capital: $255.8M', style='List Bullet')
doc.add_paragraph('Reserve (Emergency Fund): $12.5M (4.9%)', style='List Bullet')
doc.add_paragraph('Operating Capital (Day-to-day): $128.2M (49.8%)', style='List Bullet')
doc.add_paragraph('Investment Capital: $115.1M (44.7%)', style='List Bullet')

doc.add_heading('Returns on Capital:', 2)
returns = [
    ('RORAC', '23.1%', 'Return on Risk-Adjusted Capital - PRIMARY METRIC'),
    ('ROE', '18.2%', 'Return on Equity - profit per shareholder dollar'),
    ('ROI', '21.1%', 'Return on Investment - overall profit percentage'),
    ('ROCE', '19.7%', 'Return on Capital Employed - efficiency of capital use')
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'
hdr_cells[2].text = 'Meaning'

for metric, value, meaning in returns:
    row_cells = table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = value
    row_cells[2].text = meaning

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Efficiency Grade: ').bold = True
p.add_run('A+ EXCELLENT - Top tier performance')

doc.add_heading('Growth Forecast (Year-over-Year):', 2)
forecasts = [
    'Premium Growth: +12.0% (acquiring more customers)',
    'RORAC Growth: +8.5% (improving returns)',
    'Capital Gain: +15.2% (growing available capital)',
    'Profit Growth: +18.3% (bottom line profit increasing rapidly)'
]
for forecast in forecasts:
    doc.add_paragraph(forecast, style='List Bullet')

doc.add_page_break()

# 7. Risk-Return Optimization
doc.add_heading('7. Risk-Return Optimization Strategy', 1)

doc.add_heading('Current Asset Allocation vs. Target:', 2)
allocation = [
    ('Direct Business', '65%', '60%', '+5% above target (slightly aggressive)'),
    ('Treaty Reinsurance', '25%', '30%', '-5% below target (more room)'),
    ('Facultative', '10%', '10%', 'At target (balanced)')
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Type'
hdr_cells[1].text = 'Current'
hdr_cells[2].text = 'Target'
hdr_cells[3].text = 'Analysis'

for type_name, current, target, analysis in allocation:
    row_cells = table.add_row().cells
    row_cells[0].text = type_name
    row_cells[1].text = current
    row_cells[2].text = target
    row_cells[3].text = analysis

doc.add_paragraph()

doc.add_heading('Risk-Return Profile:', 2)
doc.add_paragraph('Expected Return: 23.5%', style='List Bullet')
doc.add_paragraph('Volatility (Risk): 8.2%', style='List Bullet')
doc.add_paragraph('Sharpe Ratio: 2.87 (excellent - high return for risk taken)', style='List Bullet')
doc.add_paragraph('Maximum Drawdown: -12.3% (worst case scenario)', style='List Bullet')
doc.add_paragraph('Risk Grade: A (Low-Moderate) - Conservative but profitable', style='List Bullet')

doc.add_heading('Optimization Opportunities:', 2)
doc.add_paragraph('Current Efficiency: 0.85 (out of 1.0)', style='List Bullet')
doc.add_paragraph('Potential Efficiency: 0.92 (possible improvement)', style='List Bullet')
doc.add_paragraph('Available Upside: +8.2% (optimization potential)', style='List Bullet')

doc.add_heading('Recommended Actions:', 2)
doc.add_paragraph('Reduce Casualty segment from 33% → 28% (-5%)', style='List Number')
doc.add_paragraph('Increase Marine segment from 12.4% → 17% (+4.6%)', style='List Number')
doc.add_paragraph('Expand Asia Pacific from 19.7% → 25% (+5.3%)', style='List Number')

doc.add_page_break()

# 8. Dashboard Benefits
doc.add_heading('8. Dashboard Benefits Summary', 1)

doc.add_heading('For Executives:', 2)
doc.add_paragraph('Real-time portfolio health monitoring', style='List Bullet')
doc.add_paragraph('Quick identification of risk areas', style='List Bullet')
doc.add_paragraph('Data-driven decision making', style='List Bullet')
doc.add_paragraph('ROI tracking and optimization opportunities', style='List Bullet')

doc.add_heading('For Portfolio Managers:', 2)
doc.add_paragraph('Segment performance tracking across LOBs', style='List Bullet')
doc.add_paragraph('Geographic profitability analysis', style='List Bullet')
doc.add_paragraph('Rebalancing recommendations', style='List Bullet')
doc.add_paragraph('Capital allocation optimization', style='List Bullet')

doc.add_heading('For Risk Managers:', 2)
doc.add_paragraph('Comprehensive risk assessment framework', style='List Bullet')
doc.add_paragraph('Early warning systems for concentration', style='List Bullet')
doc.add_paragraph('Capital adequacy monitoring', style='List Bullet')
doc.add_paragraph('Stress testing scenarios', style='List Bullet')

doc.add_heading('For Investors:', 2)
doc.add_paragraph('Profitability metrics (RORAC, ROE, ROI)', style='List Bullet')
doc.add_paragraph('Growth trajectory visibility', style='List Bullet')
doc.add_paragraph('Risk-adjusted returns clarity', style='List Bullet')
doc.add_paragraph('Dividend potential assessment', style='List Bullet')

doc.add_page_break()

# Quick Reference Summary
doc.add_heading('Quick Reference: Key Numbers Summary', 1)

summary_data = {
    'Revenue & Growth': {
        'GWP': '$256.9M (+5.2%)',
        'Premium Growth': '12.0%',
        'Total Profit': '$138M (estimated)',
    },
    'Efficiency Metrics': {
        'RORAC': '23.1%',
        'ROE': '18.2%',
        'ROI': '21.1%',
        'Efficiency Grade': 'A+ EXCELLENT'
    },
    'Risk Metrics': {
        'Capital Utilization': '98.9%',
        'Diversification Score': '0.76/1.0',
        'Claims Ratio': '58%',
        'Risk Score': '100/100'
    },
    'Market Position': {
        'Primary Market': 'North America (39.6%, $102M)',
        'Largest LOB': 'Casualty (33%, $85M)',
        'Fastest Growth': 'Asia Pacific (19.7%)',
        'Market Outlook': '📈 Positive'
    }
}

for category, metrics in summary_data.items():
    doc.add_heading(category, 2)
    for metric, value in metrics.items():
        p = doc.add_paragraph()
        p.add_run(f'{metric}: ').bold = True
        p.add_run(value)

# Save document
output_path = 'Portfolio_Dashboard_Guide.docx'
doc.save(output_path)
print(f'✅ Document created successfully: {output_path}')
print(f'📂 Location: C:\\Users\\m107\\OneDrive - Capgemini\\Desktop\\MCP\\portfolio-agent\\{output_path}')
